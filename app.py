import streamlit as st
from io import BytesIO
from docx import Document
import PyPDF2
import re
from datetime import datetime
from Bio import Align
from Bio.Seq import Seq


# ------------------------
# UGENE-style Consensus Logic with Tolerance
# ------------------------

def reverse(seq):
    return seq[::-1]

def complement(seq):
    return seq.translate(str.maketrans("ACGTacgt", "TGCAtgca"))

def reverse_complement(seq):
    return complement(reverse(seq))

def find_overlap_suffix_prefix(s1, s3, tolerance):
    """Return longest approximate match between suffix of s1 and prefix of s3,
    building matched string using forward read's base on mismatch."""
    max_len = min(len(s1), len(s3))
    for i in range(max_len, 0, -1):
        suffix = s1[-i:]
        prefix = s3[:i]
        mismatches = sum(1 for a, b in zip(suffix, prefix) if a != b)
        if mismatches <= tolerance:
            # Build a2 with forward preference
            match_chars = [
                a if a == b else a  # prefer forward read's base on mismatch
                for a, b in zip(suffix, prefix)
            ]
            return "".join(match_chars), i
    return "", 0

def ugene_style_consensus(s1, s2, tolerance=0):
    s1 = s1.strip().replace("\n", "").upper()
    s2 = s2.strip().replace("\n", "").upper()

    s2_rev = reverse(s2)
    s3 = reverse_complement(s2)

    matching_str, match_len = find_overlap_suffix_prefix(s1, s3, tolerance)

    if match_len == 0:
        return None, s1, s2, s2_rev, s3, "", ""

    a1 = s1[:-match_len].lower()
    a2 = matching_str.upper()
    unmatched_s3 = s3[match_len:]
    a3 = reverse_complement(unmatched_s3).lower()

    consensus = a1 + a2 + a3
    return consensus, s1, s2, s2_rev, s3, a2, a3

def parse_fasta(text_data):
    sequences = {}
    current_label = None
    for line in text_data.strip().splitlines():
        line = line.strip()
        if line.startswith(">"):
            current_label = line[1:]
            sequences[current_label] = ""
        elif current_label:
            sequences[current_label] += line.upper()
    return sequences


# ------------------------
# Biopython Consensus Logic
# ------------------------

def biopython_consensus(forward, reverse):
    """Generate consensus using BioPython alignment with UGENE-style merge rules."""
    forward = forward.strip().replace("\n", "").upper()
    reverse = reverse.strip().replace("\n", "").upper()
    reverse_complement_seq = str(Seq(reverse).reverse_complement())

    # Global alignment so both full sequences are considered
    aligner = Align.PairwiseAligner()
    aligner.mode = 'global'
    aligner.match_score = 2       # Match reward
    aligner.mismatch_score = -1   # Mismatch penalty
    aligner.open_gap_score = -2
    aligner.extend_gap_score = -0.5

    alignments = aligner.align(forward, reverse_complement_seq)
    if not alignments:
        return None, forward, reverse_complement_seq

    # Take the best alignment
    best_alignment = alignments[0]
    seq1 = str(best_alignment[0])  # aligned forward
    seq2 = str(best_alignment[1])  # aligned reverse complement

    consensus = []
    for base1, base2 in zip(seq1, seq2):
        if base1 == base2:
            consensus.append(base1)
        elif base1 == "-":  # gap in forward → take reverse base
            consensus.append(base2)
        elif base2 == "-":  # gap in reverse → take forward base
            consensus.append(base1)
        else:
            # UGENE-style: prefer forward read on mismatch
            consensus.append(base1)

    # Remove gaps from final consensus
    consensus_seq = "".join(consensus).replace("-", "")

    return consensus_seq, forward, reverse_complement_seq


# ------------------------
# IGHV PDF extraction
# ------------------------

def extract_from_pdf(pdf_bytes):
    pdf_reader = PyPDF2.PdfReader(BytesIO(pdf_bytes))
    text = "".join(page.extract_text() or "" for page in pdf_reader.pages)

    gene_names_list = []
    percent_identity_str = None
    ratio_str = None

    pattern_section = re.compile(r"Sequences\s+pr.*?alignmen", re.IGNORECASE | re.DOTALL)
    section_match = pattern_section.search(text)
    if section_match:
        start_pos = section_match.end()
        lines_after_section = text[start_pos:].strip().splitlines()
        for line in lines_after_section:
            for word in line.strip().split():
                if word.startswith("IGHV"):
                    gene_names_list.append(word)
            if gene_names_list:
                break

    pattern_total = r"Total\s+(\d+)\s+(\d+)\s+\d+\s+\d+\s+([\d\.]+)"
    match_total = re.search(pattern_total, text)
    if match_total:
        length = int(match_total.group(1))
        matches = int(match_total.group(2))
        identity = float(match_total.group(3))
        percent_identity_str = f"{identity}%"
        ratio_str = f"{matches}/{length}"

    return gene_names_list, percent_identity_str, ratio_str


# ------------------------
# DOCX Report Generator
# ------------------------

def generate_docx_report(gene_names_list, percent_identity_str, ratio_str, template_bytes, sample_id_text):
    gene_name_str = ", ".join(str(g).strip() for g in gene_names_list)
    percent_identity = float(percent_identity_str.strip('%'))
    mutation_percent = round(100 - percent_identity, 1)
    mutation_percent_str = f"{mutation_percent}%"

    if any("3-21" in g for g in gene_names_list):
        prognosis_text = (
            "BAD\n"
            "Note: CLL expressing the IGHV 3-21 variable region gene segment "
            "have a poorer prognosis regardless of IGHV mutation status."
        )
        prognosis_single_line = prognosis_text
    else:
        if mutation_percent <= 2.0:
            prognosis_text = "BAD"
        elif 2.1 <= mutation_percent <= 3.0:
            prognosis_text = "Borderline with intermediate clinical course"
        else:
            prognosis_text = "GOOD"
        prognosis_single_line = prognosis_text

    doc = Document(BytesIO(template_bytes))

    for section in doc.sections:
        for para in section.header.paragraphs:
            if "Sample ID" in para.text:
                para.text = f"Sample ID: {sample_id_text}"

    for para in doc.paragraphs:
        if "Sample ID" in para.text:
            para.text = f"Sample ID: {sample_id_text}"
        elif "Mutation detected:" in para.text:
            para.text = f"Mutation detected: {mutation_percent_str}"
        elif "Clinical Prognosis" in para.text:
            idx = doc.paragraphs.index(para)
            if idx + 1 < len(doc.paragraphs):
                doc.paragraphs[idx + 1].text = prognosis_text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "Sample ID" in cell.text:
                    cell.text = f"Sample ID: {sample_id_text}"
                elif "Mutation detected:" in cell.text:
                    row.cells[1].text = mutation_percent_str
                elif "Clinical Prognosis:" in cell.text:
                    row.cells[1].text = prognosis_single_line

    for table in doc.tables:
        first_row = [cell.text.strip() for cell in table.rows[0].cells]
        if "Name" in first_row and "Percentage Identity Detected" in first_row:
            for i in range(len(table.rows) - 1, 0, -1):
                table._tbl.remove(table.rows[i]._tr)
            row_cells = table.add_row().cells
            row_cells[0].text = gene_name_str
            row_cells[1].text = percent_identity_str
            row_cells[2].text = mutation_percent_str
            row_cells[3].text = ratio_str
            break

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


# ------------------------
# STREAMLIT APP
# ------------------------

st.set_page_config(page_title="IGHV Report Generator", layout="centered")
st.title("IGHV Report Generator")
st.caption("*A Wobble Base Bioresearch proprietary software*")

# ------------------------
# SECTION 1: Consensus Generator
# ------------------------

st.header("1. Consensus Generator")
cap3_input_file = st.file_uploader("Upload FASTA File with Forward (_F) and Reverse (_R) Reads", type=["txt", "fasta"])

# 👇 New: Select consensus method
consensus_method = st.radio("Choose Consensus Method", options=["Contig generation IM1", "Contig generation IM2"])

# 👇 Tolerance input (UGENE only)
if consensus_method == "Contig generation IM1":
    tolerance_limit = st.number_input("Mismatch Tolerance (Max mismatches allowed in overlap)", min_value=0, max_value=100, value=0, step=1)

if cap3_input_file:
    content = cap3_input_file.read().decode("utf-8")
    sequences = parse_fasta(content)
    forward = next((v for k, v in sequences.items() if k.endswith("_F")), None)
    reverse_seq = next((v for k, v in sequences.items() if k.endswith("_R")), None)

    if forward and reverse_seq:
        if consensus_method == "Contig generation IM1":
            consensus, s1, s2, s2_rev, s3, a2, a3 = ugene_style_consensus(forward, reverse_seq, tolerance=tolerance_limit)
        else:
            consensus, s1, s2 = biopython_consensus(forward, reverse_seq)

        st.subheader("Inputs and Intermediates")
        st.text_area("Forward Read (s1)", s1, height=100)
        st.text_area("Reverse Read (s2)", s2, height=100)

        if consensus:
            st.success(f"{consensus_method} Consensus Generated")
            st.code(consensus, language="text")

            consensus_bytes = BytesIO(consensus.encode("utf-8"))
            st.download_button(
                label="Download Consensus (.txt)",
                data=consensus_bytes,
                file_name="consensus_output.txt",
                mime="text/plain"
            )
        else:
            st.warning("No match found. Manual intervention needed.")
    else:
        st.error("Could not find both _F and _R sequences in file.")

# ------------------------
# SECTION 2: IGHV Report Generator
# ------------------------

st.header("2. IGHV DOCX Report Generator")
st.markdown("[Downstream analysis](https://www.ncbi.nlm.nih.gov/igblast/)")

pdf_file = st.file_uploader("PDF file (IgBLAST results)", type="pdf")
docx_template_file = st.file_uploader("Word Template (.docx)", type="docx")
ab1_file = st.file_uploader(".ab1 Sequence File", type="ab1")

if st.button("Generate IGHV Report"):
    if pdf_file and docx_template_file and ab1_file:
        ab1_filename = ab1_file.name
        sample_id = ab1_filename.split("(")[0].strip()
        date_str = datetime.now().strftime("%d-%m-%Y")
        final_name = f"{sample_id} ({date_str})"
        output_filename = f"{final_name}.docx"

        st.info(f"Sample ID: {sample_id}")
        st.info(f"Output file name: {output_filename}")

        pdf_bytes = pdf_file.read()
        gene_names_list, percent_identity_str, ratio_str = extract_from_pdf(pdf_bytes)

        st.write("Gene Names:", gene_names_list)
        st.write("Percent Identity:", percent_identity_str)
        st.write("Ratio:", ratio_str)

        if gene_names_list and percent_identity_str and ratio_str:
            docx_bytes = generate_docx_report(
                gene_names_list=gene_names_list,
                percent_identity_str=percent_identity_str,
                ratio_str=ratio_str,
                template_bytes=docx_template_file.read(),
                sample_id_text=final_name
            )

            st.download_button(
                label="Download Report",
                data=docx_bytes,
                file_name=output_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.error("Could not extract all required fields from PDF.")
    else:
        st.warning("Please upload all three files before generating the report.")
